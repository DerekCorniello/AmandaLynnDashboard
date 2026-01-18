from django.http import JsonResponse
from django.views import View
from django.db import connection
from django.db.models import Q, Sum
from django.db.models.functions import TruncDate
from .models import Product, Expense, Transaction
from collections import defaultdict
from datetime import datetime
import json


def apply_sorting_and_filtering(queryset, request, allowed_sort_fields):
    filters = {key: value for key, value in request.GET.items() if key not in [
        'sort_by', 'order', 'search', 'show_retired']}

    # Apply filters from the request
    queryset = queryset.filter(**filters)

    # Apply search functionality
    search_query = request.GET.get('search', '')
    if search_query:
        search_conditions = Q()  # An empty Q object for OR conditions
        for field in allowed_sort_fields:
            search_conditions |= Q(**{f'{field}__icontains': search_query})
        queryset = queryset.filter(search_conditions)

    # Apply sorting
    sort_by = request.GET.get('sort_by')
    order = request.GET.get('order', 'asc')

    if sort_by in allowed_sort_fields:
        if order == 'desc':
            sort_by = f'-{sort_by}'
        queryset = queryset.order_by(sort_by)

    return queryset


class GraphData(View):

    def _get_money_data(self, timescale):
        today = datetime.now()
        match timescale:
            case 'week':
                start_date = today - datetime.timedelta(days=7)
            case 'month':
                start_date = today - datetime.timedelta(days=30)
            case '3month':
                start_date = today - datetime.timedelta(days=90)
            case '6month':
                start_date = today - datetime.timedelta(days=180)
            case 'year':
                start_date = today - datetime.timedelta(days=365)
            case 'all':
                start_date = datetime.min
            case _:
                return JsonResponse(
                    {'error': 'Invalid time scale'},
                    status=400)

        # Fetch data from Expense model and group by date
        expenses_query = (
            Expense.objects
            .filter(date__gte=start_date)
            .annotate(date=TruncDate('date'))
            .values('date')
            .annotate(total_expense=Sum('price'))
            .values('date', 'total_expense')
        )

        expenses = list(expenses_query)

        # Fetch data from Transaction model and group by date
        transactions_query = (
            Transaction.objects
            .filter(date__gte=start_date)
            .annotate(date=TruncDate('date'))
            .values('date')
            .annotate(total_income=Sum('total'))
            .values('date', 'total_income')
        )

        transactions = list(transactions_query)

        # Aggregate daily data into dictionaries
        daily_expenses = defaultdict(float)
        daily_income = defaultdict(float)

        for expense in expenses:
            daily_expenses[expense['date']] = expense['total_expense']

        for transaction in transactions:
            daily_income[transaction['date']] = transaction['total_income']

        # Prepare data for chart
        dates = sorted(
            set(daily_expenses.keys()).union(daily_income.keys()))
        expense_data = [daily_expenses.get(date, 0) for date in dates]
        income_data = [daily_income.get(date, 0) for date in dates]
        revenue_data = [income - expense for income,
                        expense in zip(income_data, expense_data)]

        return {
            'labels': dates,
            'datasets': [
                {
                    'label': 'Total Income',
                    'data': income_data,
                    'borderColor': '#42A5F5',
                    'backgroundColor': 'rgba(66, 165, 245, 0.2)'
                },
                {
                    'label': 'Total Spending',
                    'data': expense_data,
                    'borderColor': '#66BB6A',
                    'backgroundColor': 'rgba(102, 187, 106, 0.2)'
                },
                {
                    'label': 'Revenue',
                    'data': revenue_data,
                    'borderColor': '#FF7043',
                    'backgroundColor': 'rgba(255, 112, 67, 0.2)'
                }
            ]
        }

    def _get_product_data(self, timescale):
        return

    def _get_timeseries_data(self, request_data):
        try:
            years_str = request_data.get('years', str(datetime.now().year))
            metrics_str = request_data.get('metrics', 'revenue,profit')
            products_str = request_data.get('products', 'all')

            years = [int(y.strip()) for y in years_str.split(',')]
            metrics = [m.strip().lower() for m in metrics_str.split(',')]
            products = [p.strip() for p in products_str.split(',')] if products_str.lower() != 'all' else None

            all_datasets = []
            all_labels = []

            colors = ['#42A5F5', '#FF6384', '#4BC0C0', '#FFCE56', '#36A2EB', '#9966FF', '#FF9F40']

            year_colors = {}
            for i, year in enumerate(years):
                year_colors[year] = colors[i % len(colors)]

            for year in years:
                year_start = datetime(year, 1, 1)
                year_end = datetime(year, 12, 31)

                month_labels = [f'{year}-{str(m).zfill(2)}' for m in range(1, 13)]

                revenue_by_month = defaultdict(float)
                loss_by_month = defaultdict(float)

                transactions = Transaction.objects.filter(date__range=(year_start, year_end))
                expenses = Expense.objects.filter(date__range=(year_start, year_end))

                for transaction in transactions:
                    month = transaction.date.month
                    revenue_by_month[month] += float(transaction.total)

                for expense in expenses:
                    month = expense.date.month
                    loss_by_month[month] += float(expense.price)

                if 'revenue' in metrics:
                    all_datasets.append({
                        'label': f'Revenue {year}',
                        'data': [revenue_by_month[m] for m in range(1, 13)],
                        'borderColor': year_colors[year],
                        'backgroundColor': 'transparent',
                        'borderWidth': 2,
                        'pointRadius': 4,
                        'fill': False
                    })

                if 'loss' in metrics:
                    all_datasets.append({
                        'label': f'Loss {year}',
                        'data': [loss_by_month[m] for m in range(1, 13)],
                        'borderColor': year_colors[year],
                        'backgroundColor': 'transparent',
                        'borderWidth': 2,
                        'pointRadius': 4,
                        'fill': False
                    })

                if 'profit' in metrics:
                    all_datasets.append({
                        'label': f'Profit {year}',
                        'data': [revenue_by_month[m] - loss_by_month[m] for m in range(1, 13)],
                        'borderColor': year_colors[year],
                        'backgroundColor': 'transparent',
                        'borderWidth': 2,
                        'pointRadius': 4,
                        'fill': False
                    })

                all_labels.extend(month_labels)

            if 'product_sales' in metrics:
                selected_products = products
                if not selected_products:
                    all_products = Product.objects.filter(is_retired=False)
                    selected_products = [p.name for p in all_products]

                for year in years:
                    year_start = datetime(year, 1, 1)
                    year_end = datetime(year, 12, 31)

                    product_sales = {}
                    for product in selected_products:
                        product_sales[product] = defaultdict(int)

                    transactions = Transaction.objects.filter(date__range=(year_start, year_end))

                    for transaction in transactions:
                        if isinstance(transaction.products, str):
                            products_list = [p.strip().strip("[]'") for p in transaction.products.split(',')]
                        else:
                            products_list = list(transaction.products)

                        for product_name in products_list:
                            if product_name in product_sales:
                                product_sales[product_name][transaction.date.month] += 1

                    month_labels = [f'{year}-{str(m).zfill(2)}' for m in range(1, 13)]

                    for i, product_name in enumerate(selected_products):
                        all_labels.extend(month_labels)
                        color = colors[(i + len(years)) % len(colors)]
                        all_datasets.append({
                            'label': f'{product_name} Sales {year}',
                            'data': [product_sales[product_name][m] for m in range(1, 13)],
                            'borderColor': color,
                            'backgroundColor': 'transparent',
                            'borderWidth': 2,
                            'pointRadius': 4,
                            'fill': False
                        })

            if not all_datasets:
                return {
                    'labels': [],
                    'datasets': []
                }

            return {
                'labels': list(dict.fromkeys(all_labels)),
                'datasets': all_datasets
            }

        except Exception as e:
            print(f'Error in _get_timeseries_data: {e}')
            return {
                'labels': [],
                'datasets': []
            }

    def post(self, request):
        try:
            data = json.loads(request.body)
            if data['graph'] == 'money':
                res = self._get_money_data(data['timescale'])
            elif data['graph'] == 'product':
                res = self._get_product_data(data['timescale'])
            elif data['graph'] == 'timeseries':
                res = self._get_timeseries_data(data)
            else:
                res = None
                raise KeyError
            return JsonResponse(res, safe=False)
        except KeyError as e:
            return JsonResponse({'error': f'Graph requested `{data["graph"]}` is not available: {e}'}, status=404)
        except Exception as e:
            return JsonResponse({'error': f'Internal Server Error:{e}'}, status=500)


class Status(View):
    def get(self, request):
        try:
            connection.ensure_connection()
            return JsonResponse({"status": "ok"}, status=200)
        except Exception:
            return JsonResponse({"status": "error"}, status=500)


class ProductList(View):
    def get(self, request):
        try:
            show_retired = request.GET.get(
                "show_retired", "false").lower() == "true"
            products = Product.objects.all()

            if not show_retired:
                products = products.filter(is_retired=False)

            allowed_sort_fields = ['name', 'price', 'stock', 'number_sold']
            products = apply_sorting_and_filtering(
                products, request, allowed_sort_fields)

            products = list(products.values())
            return JsonResponse(products, safe=False)
        except Exception as e:
            return JsonResponse({'error': f'Internal Server Error: {e}'}, status=500)


class ProductDelete(View):
    def delete(self, request, pk):
        try:
            product = Product.objects.get(pk=pk)
            product.delete()
            return JsonResponse({'status': 'success'}, status=204)
        except Product.DoesNotExist:
            return JsonResponse({'error': 'Not found'}, status=404)
        except Exception as e:
            return JsonResponse({'error': f'Internal Server Error:{e}'}, status=500)


class ProductCreate(View):
    def post(self, request):
        try:
            data = json.loads(request.body)
            valid_products = Product.objects.values_list('name', flat=True)
            if data['name'].lower() in [p.lower() for p in valid_products]:
                return JsonResponse({'error': f'Product already exists: {data["name"]}'}, status=400)
            if data['name'].lower() == 'unknown':
                return JsonResponse({'error': 'Cannot create Unknown Product'}, status=400)

            product = Product.objects.create(
                name=data['name'],
                stock=data['stock'],
                price=data['price'],
                number_sold=data['number_sold']
            )
            return JsonResponse({
                'id': product.id,
                'name': product.name,
                'stock': product.stock,
                'price': product.price,
                'number_sold': product.number_sold
            }, status=201)
        except KeyError as e:
            return JsonResponse({'error': f'Missing field: {str(e)}'}, status=400)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)
        except Exception as e:
            return JsonResponse({'error': f'Internal Server Error:{e}'}, status=500)


class ProductUpdate(View):
    def put(self, request, pk):
        try:
            data = json.loads(request.body)
            product = Product.objects.get(pk=pk)
            valid_products = Product.objects.values_list('name', flat=True)

            if data['name'].lower() in [p.lower() for p in valid_products] and data['name'].lower() != product.name.lower():
                return JsonResponse({'error': f'Product already exists: {data["name"]}'}, status=400)
            if data['name'].lower() == 'unknown':
                return JsonResponse({'error': 'Cannot create Unknown Product'}, status=400)

            for field in ['name', 'stock', 'price', 'number_sold', 'is_retired']:
                if field in data:
                    setattr(product, field, data[field])

            product.save()
            return JsonResponse({
                'id': product.id,
                'name': product.name,
                'stock': product.stock,
                'price': product.price,
                'number_sold': product.number_sold
            }, status=205)
        except Product.DoesNotExist:
            return JsonResponse({'error': 'Not found'}, status=404)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)
        except Exception as e:
            return JsonResponse({'error': f'Internal Server Error:{e}'}, status=500)


class ExpenseList(View):
    def get(self, request):
        try:
            expenses = Expense.objects.all()
            allowed_sort_fields = ['name', 'date', 'price', 'type']
            expenses = apply_sorting_and_filtering(
                expenses, request, allowed_sort_fields)
            expenses = list(expenses.values())
            return JsonResponse(expenses, safe=False)
        except Exception as e:
            return JsonResponse({'error': f'Internal Server Error:{e}'}, status=500)


class ExpenseDelete(View):
    def delete(self, request, pk):
        try:
            expense = Expense.objects.get(pk=pk)
            expense.delete()
            return JsonResponse({'status': 'success'}, status=204)
        except Expense.DoesNotExist:
            return JsonResponse({'error': 'Not found'}, status=404)
        except Exception as e:
            return JsonResponse({'error': f'Internal Server Error:{e}'}, status=500)


class ExpenseCreate(View):
    def post(self, request):
        try:
            data = json.loads(request.body)
            expense = Expense.objects.create(
                name=data['name'],
                date=data['date'],
                type=data['type'],
                price=data['price']
            )
            return JsonResponse({
                'id': expense.id,
                'name': expense.name,
                'date': expense.date,
                'type': expense.type,
                'price': expense.price
            }, status=201)
        except KeyError as e:
            return JsonResponse({'error': f'Missing field: {str(e)}'}, status=400)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)
        except Exception as e:
            return JsonResponse({'error': f'Internal Server Error:{e}'}, status=500)


class ExpenseUpdate(View):
    def put(self, request, pk):
        try:
            data = json.loads(request.body)
            expense = Expense.objects.get(pk=pk)

            for field in ['name', 'date', 'type', 'price']:
                if field in data:
                    setattr(expense, field, data[field])

            expense.save()
            return JsonResponse({
                'date': expense.date,
                'name': expense.name,
                'type': expense.type,
                'price': expense.price,
            }, status=205)
        except Expense.DoesNotExist:
            return JsonResponse({'error': 'Not found'}, status=404)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)
        except Exception as e:
            return JsonResponse({'error': f'Internal Server Error:{e}'}, status=500)


class TransactionList(View):
    def get(self, request):
        try:
            transactions = Transaction.objects.all()
            allowed_sort_fields = ['date', 'total', 'type']
            transactions = apply_sorting_and_filtering(
                transactions, request, allowed_sort_fields)

            transaction_data = []
            for transaction in transactions:
                # if this is a string, we do not want to use the comma join
                # or else it will list-ify the string, and list its chars
                if isinstance(transaction.products, str):
                    products = transaction.products
                else:
                    products = ', '.join(transaction.products)
                # trims the '' and [] off of the string
                products = products.replace("'", "")[1:-1]
                transaction_data.append({
                    'id': transaction.id,
                    'total': transaction.total,
                    'date': transaction.date,
                    'type': transaction.type,
                    'products': products
                })

            return JsonResponse(transaction_data, safe=False)
        except Exception as e:
            return JsonResponse({'error': f'Internal Server Error:{e}'}, status=500)


class TransactionDelete(View):
    def delete(self, request, pk):
        try:
            transaction = Transaction.objects.get(pk=pk)
            transaction.delete()
            return JsonResponse({'status': 'success'}, status=204)
        except Transaction.DoesNotExist:
            return JsonResponse({'error': 'Not found'}, status=404)
        except Exception as e:
            return JsonResponse({'error': f'Internal Server Error:{e}'}, status=500)


class TransactionCreate(View):
    def post(self, request):
        try:
            data = json.loads(request.body)

            product_names_list = data.get('products', [])
            if isinstance(product_names_list, str):
                product_names_list = [p.strip()
                                      for p in product_names_list.split(',')]
            valid_products = Product.objects.values_list('name', flat=True)
            print(valid_products)
            for product in product_names_list:
                if product.lower() not in [p.lower().strip() for p in valid_products] and product.lower() != 'unknown':
                    print(f'Product does not exist: {product}')
                    return JsonResponse({'error': f'Product does not exist: {product}'}, status=400)

            transaction = Transaction.objects.create(
                total=data['total'],
                date=data['date'],
                type=data['type']
            )

            transaction.products = product_names_list
            transaction.save()

            return JsonResponse({
                'id': transaction.id,
                'total': transaction.total,
                'date': transaction.date,
                'type': transaction.type,
                'products': product_names_list
            }, status=201)

        except KeyError as e:
            return JsonResponse({'error': f'Missing field: {str(e)}'}, status=400)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)
        except Exception as e:
            return JsonResponse({'error': f'Internal Server Error:{e}'}, status=500)


class TransactionUpdate(View):
    def put(self, request, pk):
        try:
            data = json.loads(request.body)
            transaction = Transaction.objects.get(pk=pk)

            for field in ['total', 'date', 'type']:
                if field in data:
                    setattr(transaction, field, data[field])

            product_names_list = data.get('products', [])
            if isinstance(product_names_list, str):
                product_names_list = [p.strip()
                                      for p in product_names_list.split(',')]

            valid_products = Product.objects.values_list('name', flat=True)

            for product in product_names_list:
                if product.lower() not in [p.lower().strip() for p in valid_products] and product.lower() != 'unknown':
                    return JsonResponse({'error': f'Product does not exist: {product}'}, status=400)

            transaction.products = product_names_list
            transaction.save()

            return JsonResponse({
                'id': transaction.id,
                'total': transaction.total,
                'date': transaction.date,
                'type': transaction.type,
                'products': product_names_list
            }, status=205)

        except Transaction.DoesNotExist:
            return JsonResponse({'error': 'Not found'}, status=404)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)
        except Exception as e:
            return JsonResponse({'error': f'Internal Server Error: {e}'}, status=500)


class ProductComparison(View):
    def get(self, request):
        try:
            products = Product.objects.filter(is_retired=False)

            product_details = []
            stock_data = []
            number_sold_data = []

            for product in products:
                product_details.append({
                    'name': product.name,
                    'price': float(product.price)
                })
                stock_data.append(product.stock)
                number_sold_data.append(product.number_sold)

            return JsonResponse({
                'labels': [p['name'] for p in product_details],
                'datasets': [
                    {
                        'label': 'Stock',
                        'data': stock_data,
                        'backgroundColor': 'rgba(54, 162, 235, 0.7)',
                        'borderColor': 'rgba(54, 162, 235, 1)',
                        'borderWidth': 1
                    },
                    {
                        'label': 'Number Sold',
                        'data': number_sold_data,
                        'backgroundColor': 'rgba(255, 99, 132, 0.7)',
                        'borderColor': 'rgba(255, 99, 132, 1)',
                        'borderWidth': 1
                    }
                ],
                'product_details': product_details
            })
        except Exception as e:
            return JsonResponse({'error': f'Internal Server Error: {e}'}, status=500)


class SaveData(View):
    def post(self, request):
        try:
            # Force database commit (Django handles this automatically, but we'll provide feedback)
            from django.db import transaction
            transaction.commit()

            # Return success with current timestamp
            import datetime
            return JsonResponse({
                "message": "Data saved successfully",
                "timestamp": datetime.datetime.now().isoformat(),
                "status": "saved"
            }, status=200)
        except Exception as e:
            return JsonResponse({
                "message": f"Save failed: {str(e)}",
                "status": "error"
            }, status=500)


class HomeView(View):
    def get(self, request):
        try:
            import os
            from django.conf import settings

            index_path = os.path.join(settings.BASE_DIR, 'frontend', 'dist', 'index.html')

            if os.path.exists(index_path):
                with open(index_path, 'r', encoding='utf-8') as f:
                    content = f.read()

                # Add proper content-type for HTML
                from django.http import HttpResponse
                return HttpResponse(content, content_type='text/html')
            else:
                from django.http import HttpResponseNotFound
                return HttpResponseNotFound('Frontend not found. Please build the Vue app.')
        except Exception as e:
            from django.http import HttpResponseServerError
            return HttpResponseServerError(f'Error serving frontend: {str(e)}')


class ExportData(View):
    def get(self, request):
        try:
            data_type = request.GET.get('type', 'all')
            format_type = request.GET.get('format', 'txt')

            data = {}
            title = ''

            if data_type in ['products', 'all']:
                products = list(Product.objects.all().values())
                data['products'] = products
                title = 'Products'

            if data_type in ['expenses', 'all']:
                expenses = list(Expense.objects.all().values())
                data['expenses'] = expenses
                title = 'Expenses'

            if data_type in ['transactions', 'all']:
                transactions = list(Transaction.objects.all().values())
                data['transactions'] = transactions
                title = 'Transactions'

            # Generate filename with timestamp
            from datetime import datetime
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'amanda_lynn_{data_type}_{timestamp}'

            if format_type == 'pdf':
                return self._generate_pdf(data, filename, data_type)
            elif format_type == 'docx':
                return self._generate_docx(data, filename, data_type)
            else:
                return self._generate_txt(data, filename, data_type)

        except Exception as e:
            return JsonResponse({'error': f'Export failed: {str(e)}'}, status=500)

    def _generate_txt(self, data, filename, data_type):
        from datetime import datetime
        content = []
        content.append('=' * 60)
        content.append('AMANDA LYNN DATA EXPORT')
        content.append(f'Export Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        content.append(f'Data Type: {data_type.upper()}')
        content.append('=' * 60)
        content.append('')

        if 'products' in data and data['products']:
            content.append('-' * 40)
            content.append('PRODUCTS')
            content.append('-' * 40)
            for i, p in enumerate(data['products'], 1):
                content.append(f'{i}. {p.get("name", "N/A")}')
                content.append(f'   Stock: {p.get("stock", 0)}')
                content.append(f'   Price: ${p.get("price", 0)}')
                content.append(f'   Sold: {p.get("number_sold", 0)}')
                content.append(f'   Status: {"Active" if not p.get("is_retired") else "Retired"}')
                content.append('')

        if 'expenses' in data and data['expenses']:
            content.append('')
            content.append('-' * 40)
            content.append('EXPENSES')
            content.append('-' * 40)
            for i, e in enumerate(data['expenses'], 1):
                content.append(f'{i}. {e.get("name", "N/A")}')
                content.append(f'   Date: {e.get("date", "N/A")}')
                content.append(f'   Type: {e.get("type", "N/A")}')
                content.append(f'   Amount: ${e.get("price", 0)}')
                content.append('')

        if 'transactions' in data and data['transactions']:
            content.append('')
            content.append('-' * 40)
            content.append('TRANSACTIONS')
            content.append('-' * 40)
            for i, t in enumerate(data['transactions'], 1):
                content.append(f'{i}. Transaction #{t.get("id", "N/A")}')
                content.append(f'   Date: {t.get("date", "N/A")}')
                content.append(f'   Type: {t.get("type", "N/A")}')
                content.append(f'   Total: ${t.get("total", 0)}')
                content.append(f'   Products: {t.get("products", "N/A")}')
                content.append('')

        content.append('')
        content.append('=' * 60)
        content.append('END OF EXPORT')
        content.append('=' * 60)

        from django.http import HttpResponse
        response = HttpResponse('\n'.join(content), content_type='text/plain')
        response['Content-Disposition'] = f'attachment; filename="{filename}.txt"'
        return response

    def _generate_pdf(self, data, filename, data_type):
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors
        from io import BytesIO
        from datetime import datetime

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []

        # Title
        title_style = styles['Title']
        elements.append(Paragraph('AMANDA LYNN DATA EXPORT', title_style))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(f'Export Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', styles['Normal']))
        elements.append(Paragraph(f'Data Type: {data_type.upper()}', styles['Normal']))
        elements.append(Spacer(1, 20))

        if 'products' in data and data['products']:
            elements.append(Paragraph('PRODUCTS', styles['Heading2']))
            product_data = [['Name', 'Stock', 'Price', 'Sold', 'Status']]
            for p in data['products']:
                product_data.append([
                    p.get('name', 'N/A'),
                    str(p.get('stock', 0)),
                    f'${p.get("price", 0)}',
                    str(p.get('number_sold', 0)),
                    'Active' if not p.get('is_retired') else 'Retired'
                ])
            t = Table(product_data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 20))

        if 'expenses' in data and data['expenses']:
            elements.append(Paragraph('EXPENSES', styles['Heading2']))
            expense_data = [['Name', 'Date', 'Type', 'Amount']]
            for e in data['expenses']:
                expense_data.append([
                    e.get('name', 'N/A'),
                    str(e.get('date', 'N/A')),
                    e.get('type', 'N/A'),
                    f'${e.get("price", 0)}'
                ])
            t = Table(expense_data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 20))

        if 'transactions' in data and data['transactions']:
            elements.append(Paragraph('TRANSACTIONS', styles['Heading2']))
            trans_data = [['ID', 'Date', 'Type', 'Total', 'Products']]
            for t in data['transactions']:
                trans_data.append([
                    str(t.get('id', 'N/A')),
                    str(t.get('date', 'N/A')),
                    t.get('type', 'N/A'),
                    f'${t.get("total", 0)}',
                    t.get('products', 'N/A')[:30] + '...' if len(str(t.get('products', ''))) > 30 else t.get('products', 'N/A')
                ])
            t = Table(trans_data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            elements.append(t)

        doc.build(elements)

        from django.http import HttpResponse
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
        return response

    def _generate_docx(self, data, filename, data_type):
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading('AMANDA LYNN DATA EXPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'Export Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Data Type: {data_type.upper()}')
        doc.add_paragraph('')

        if 'products' in data and data['products']:
            doc.add_heading('PRODUCTS', level=1)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Name'
            hdr_cells[1].text = 'Stock'
            hdr_cells[2].text = 'Price'
            hdr_cells[3].text = 'Sold'
            hdr_cells[4].text = 'Status'

            for p in data['products']:
                row_cells = table.add_row().cells
                row_cells[0].text = p.get('name', 'N/A')
                row_cells[1].text = str(p.get('stock', 0))
                row_cells[2].text = f'${p.get("price", 0)}'
                row_cells[3].text = str(p.get('number_sold', 0))
                row_cells[4].text = 'Active' if not p.get('is_retired') else 'Retired'

            doc.add_paragraph('')

        if 'expenses' in data and data['expenses']:
            doc.add_heading('EXPENSES', level=1)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Name'
            hdr_cells[1].text = 'Date'
            hdr_cells[2].text = 'Type'
            hdr_cells[3].text = 'Amount'

            for e in data['expenses']:
                row_cells = table.add_row().cells
                row_cells[0].text = e.get('name', 'N/A')
                row_cells[1].text = str(e.get('date', 'N/A'))
                row_cells[2].text = e.get('type', 'N/A')
                row_cells[3].text = f'${e.get("price", 0)}'

            doc.add_paragraph('')

        if 'transactions' in data and data['transactions']:
            doc.add_heading('TRANSACTIONS', level=1)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Date'
            hdr_cells[2].text = 'Type'
            hdr_cells[3].text = 'Total'
            hdr_cells[4].text = 'Products'

            for t in data['transactions']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(t.get('id', 'N/A'))
                row_cells[1].text = str(t.get('date', 'N/A'))
                row_cells[2].text = t.get('type', 'N/A')
                row_cells[3].text = f'${t.get("total", 0)}'
                row_cells[4].text = t.get('products', 'N/A')[:30] + '...' if len(str(t.get('products', ''))) > 30 else t.get('products', 'N/A')

        # Save to buffer
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        from django.http import HttpResponse
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        return response
